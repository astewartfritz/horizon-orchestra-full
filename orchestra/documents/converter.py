"""Horizon Orchestra — Document Conversion (Pandoc).

Convert between document formats using Pandoc CLI with Python-only
fallbacks.  Supports Markdown→PDF, Markdown→DOCX, HTML→PDF, DOCX→PDF,
CSV→XLSX, and more.

Usage::

    from orchestra.documents.converter import DocumentConverter

    conv = DocumentConverter()
    output = await conv.convert("report.md", "pdf")
    html_pdf = await conv.html_to_pdf("<h1>Hello</h1>")
"""

from __future__ import annotations

import asyncio
import csv
import io
import logging
import os
import shutil
import tempfile
import uuid
from pathlib import Path
from typing import Any, Optional, Sequence

__all__ = [
    "DocumentConverter",
    "PandocNotFoundError",
]

log = logging.getLogger("orchestra.documents.converter")

_WORKSPACE = Path(os.environ.get("ORCHESTRA_WORKSPACE", "/tmp/orchestra_docs"))

# Optional dependency: markdown
try:
    import markdown as _markdown_mod
    _HAS_MARKDOWN = True
except ImportError:
    _markdown_mod = None  # type: ignore[assignment]
    _HAS_MARKDOWN = False


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------

class PandocNotFoundError(RuntimeError):
    """Raised when pandoc is not installed."""


class ConversionError(RuntimeError):
    """Raised when document conversion fails."""


# ---------------------------------------------------------------------------
# Format mappings
# ---------------------------------------------------------------------------

_PANDOC_FORMAT_MAP = {
    "md": "markdown",
    "markdown": "markdown",
    "html": "html",
    "htm": "html",
    "docx": "docx",
    "pdf": "pdf",
    "txt": "plain",
    "rst": "rst",
    "latex": "latex",
    "tex": "latex",
    "epub": "epub",
    "odt": "odt",
    "rtf": "rtf",
    "org": "org",
    "adoc": "asciidoc",
    "asciidoc": "asciidoc",
    "json": "json",
}


# ---------------------------------------------------------------------------
# DocumentConverter
# ---------------------------------------------------------------------------

class DocumentConverter:
    """Document format converter using Pandoc with Python-only fallbacks.

    Parameters
    ----------
    workspace:
        Directory for saving output files.
    pandoc_path:
        Override path to the ``pandoc`` binary.
    """

    SUPPORTED_INPUT = {"md", "markdown", "html", "htm", "docx", "txt", "rst",
                       "latex", "tex", "csv", "json", "epub", "odt", "rtf", "org"}
    SUPPORTED_OUTPUT = {"pdf", "docx", "html", "htm", "md", "markdown", "txt",
                        "rst", "latex", "tex", "epub", "odt", "rtf", "xlsx"}

    def __init__(
        self,
        workspace: str | Path | None = None,
        pandoc_path: str | None = None,
    ) -> None:
        self.workspace = Path(workspace) if workspace else _WORKSPACE / "converter"
        self.workspace.mkdir(parents=True, exist_ok=True)
        self._pandoc = pandoc_path or shutil.which("pandoc") or "pandoc"

    # ------------------------------------------------------------------
    # Guards
    # ------------------------------------------------------------------

    async def _check_pandoc(self) -> bool:
        """Check if pandoc is installed."""
        try:
            proc = await asyncio.create_subprocess_exec(
                self._pandoc, "--version",
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, _ = await proc.communicate()
            if proc.returncode == 0:
                version = stdout.decode().split("\n")[0]
                log.debug("Pandoc found: %s", version)
                return True
        except FileNotFoundError:
                        import logging as _log; _log.getLogger('documents.converter').debug('Suppressed exception', exc_info=True)
        return False

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _save_path(self, ext: str) -> Path:
        return self.workspace / f"{uuid.uuid4().hex[:12]}.{ext}"

    def _detect_format(self, path: str | Path) -> str:
        """Detect format from file extension."""
        ext = Path(path).suffix.lstrip(".").lower()
        return _PANDOC_FORMAT_MAP.get(ext, ext)

    async def _run_pandoc(
        self,
        *args: str,
    ) -> tuple[bytes, bytes]:
        """Run pandoc with the given arguments."""
        has_pandoc = await self._check_pandoc()
        if not has_pandoc:
            raise PandocNotFoundError(
                "pandoc is not installed or not found on $PATH. "
                "Install it from: https://pandoc.org/installing.html"
            )

        log.debug("Running: pandoc %s", " ".join(args))
        proc = await asyncio.create_subprocess_exec(
            self._pandoc, *args,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await proc.communicate()
        if proc.returncode != 0:
            err_msg = stderr.decode(errors="replace")[:2000]
            raise ConversionError(f"pandoc exited with code {proc.returncode}: {err_msg}")
        return stdout, stderr

    # ------------------------------------------------------------------
    # Python-only fallbacks
    # ------------------------------------------------------------------

    def _md_to_html_python(self, md: str) -> str:
        """Convert Markdown to HTML using Python."""
        if _HAS_MARKDOWN:
            return _markdown_mod.markdown(
                md,
                extensions=["tables", "fenced_code", "codehilite", "toc"],
            )
        # Minimal fallback
        import re
        html = md
        for level in range(6, 0, -1):
            pattern = r"^" + r"#" * level + r"\s+(.+)$"
            html = re.sub(pattern, rf"<h{level}>\1</h{level}>", html, flags=re.MULTILINE)
        html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
        html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)
        paragraphs = html.split("\n\n")
        processed = []
        for p in paragraphs:
            p = p.strip()
            if p and not p.startswith("<"):
                p = f"<p>{p}</p>"
            processed.append(p)
        return "\n".join(processed)

    def _csv_to_xlsx_python(self, csv_path: str | Path) -> bytes:
        """Convert CSV to XLSX using openpyxl."""
        try:
            from .xlsx import XLSXGenerator
        except ImportError:
            raise ImportError("XLSXGenerator is required for CSV→XLSX conversion.")

        with open(str(csv_path), "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = list(reader)

        gen = XLSXGenerator(workspace=self.workspace)
        return gen.from_data(rows, headers=True)

    def _html_to_pdf_python(self, html: str) -> bytes:
        """Convert HTML to PDF using Python backends."""
        try:
            from .pdf import PDFGenerator
        except ImportError:
            raise ImportError("PDFGenerator is required for HTML→PDF conversion.")

        gen = PDFGenerator(workspace=self.workspace)
        return gen.from_html(html)

    def _md_to_pdf_python(self, md: str) -> bytes:
        """Convert Markdown to PDF using Python backends."""
        try:
            from .pdf import PDFGenerator
        except ImportError:
            raise ImportError("PDFGenerator is required for MD→PDF conversion.")

        gen = PDFGenerator(workspace=self.workspace)
        return gen.from_markdown(md)

    def _md_to_docx_python(self, md: str) -> bytes:
        """Convert Markdown to DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            raise ImportError("python-docx is required for MD→DOCX fallback: pip install python-docx")

        doc = Document()

        # Simple Markdown parsing
        lines = md.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("#### "):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped[0].isdigit() and ". " in stripped[:5]:
                idx = stripped.index(". ")
                doc.add_paragraph(stripped[idx + 2:], style="List Number")
            elif stripped.startswith("> "):
                para = doc.add_paragraph(stripped[2:])
                para.style = "Quote" if "Quote" in [s.name for s in doc.styles] else para.style
            else:
                doc.add_paragraph(stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    async def convert(
        self,
        input_path: str | Path,
        output_format: str,
        *,
        output_path: str | Path | None = None,
        extra_args: Sequence[str] = (),
    ) -> Path:
        """Convert a document between formats.

        Tries Pandoc first, falls back to Python-only conversion.

        Parameters
        ----------
        input_path:
            Source document path.
        output_format:
            Target format (``pdf``, ``docx``, ``html``, ``md``, …).
        output_path:
            Destination path.  Auto-generated if *None*.
        extra_args:
            Additional arguments for Pandoc.

        Returns
        -------
        Path
            Output file path.
        """
        input_path = Path(input_path)
        out = Path(output_path) if output_path else self._save_path(output_format)
        input_ext = input_path.suffix.lstrip(".").lower()

        # Try Pandoc first
        has_pandoc = await self._check_pandoc()
        if has_pandoc:
            input_fmt = _PANDOC_FORMAT_MAP.get(input_ext, input_ext)
            output_fmt = _PANDOC_FORMAT_MAP.get(output_format, output_format)

            args = [
                str(input_path),
                "-f", input_fmt,
                "-t", output_fmt,
                "-o", str(out),
                "--standalone",
            ]
            if output_format == "pdf":
                args.extend(["--pdf-engine", "weasyprint"])
            args.extend(extra_args)

            try:
                await self._run_pandoc(*args)
                log.info("Converted %s → %s (pandoc)", input_path, out)
                return out
            except ConversionError as e:
                log.warning("Pandoc conversion failed, trying fallback: %s", e)

        # Python-only fallbacks
        content = input_path.read_text(encoding="utf-8", errors="replace")

        if input_ext in ("md", "markdown") and output_format == "pdf":
            out.write_bytes(self._md_to_pdf_python(content))
        elif input_ext in ("md", "markdown") and output_format == "docx":
            out.write_bytes(self._md_to_docx_python(content))
        elif input_ext in ("md", "markdown") and output_format in ("html", "htm"):
            html = self._md_to_html_python(content)
            out.write_text(html, encoding="utf-8")
        elif input_ext in ("html", "htm") and output_format == "pdf":
            out.write_bytes(self._html_to_pdf_python(content))
        elif input_ext == "csv" and output_format == "xlsx":
            out.write_bytes(self._csv_to_xlsx_python(input_path))
        else:
            raise ConversionError(
                f"No fallback converter for {input_ext}→{output_format}. "
                f"Install pandoc: https://pandoc.org/installing.html"
            )

        log.info("Converted %s → %s (python fallback)", input_path, out)
        return out

    async def html_to_pdf(self, html: str) -> bytes:
        """Convert HTML string to PDF.

        Parameters
        ----------
        html:
            HTML content.

        Returns
        -------
        bytes
            PDF file content.
        """
        # Try pandoc first
        has_pandoc = await self._check_pandoc()
        if has_pandoc:
            with tempfile.NamedTemporaryFile(suffix=".html", mode="w",
                                              encoding="utf-8", delete=False) as f:
                f.write(html)
                tmp_in = f.name

            tmp_out = str(self._save_path("pdf"))
            try:
                await self._run_pandoc(
                    tmp_in, "-f", "html", "-o", tmp_out,
                    "--pdf-engine", "weasyprint",
                )
                result = Path(tmp_out).read_bytes()
                return result
            except ConversionError:
                                import logging as _log; _log.getLogger('documents.converter').debug('Suppressed exception', exc_info=True)
            finally:
                Path(tmp_in).unlink(missing_ok=True)

        # Python fallback
        return self._html_to_pdf_python(html)

    async def markdown_to_docx(self, md: str) -> bytes:
        """Convert Markdown string to DOCX.

        Parameters
        ----------
        md:
            Markdown content.

        Returns
        -------
        bytes
            DOCX file content.
        """
        # Try pandoc first
        has_pandoc = await self._check_pandoc()
        if has_pandoc:
            with tempfile.NamedTemporaryFile(suffix=".md", mode="w",
                                              encoding="utf-8", delete=False) as f:
                f.write(md)
                tmp_in = f.name

            tmp_out = str(self._save_path("docx"))
            try:
                await self._run_pandoc(
                    tmp_in, "-f", "markdown", "-o", tmp_out,
                    "--standalone",
                )
                result = Path(tmp_out).read_bytes()
                return result
            except ConversionError:
                                import logging as _log; _log.getLogger('documents.converter').debug('Suppressed exception', exc_info=True)
            finally:
                Path(tmp_in).unlink(missing_ok=True)

        # Python fallback
        return self._md_to_docx_python(md)

    async def csv_to_xlsx(self, csv_path: str | Path) -> bytes:
        """Convert a CSV file to XLSX.

        Parameters
        ----------
        csv_path:
            Path to the CSV file.

        Returns
        -------
        bytes
            XLSX file content.
        """
        return self._csv_to_xlsx_python(csv_path)

    async def markdown_to_pdf(self, md: str) -> bytes:
        """Convert Markdown string to PDF.

        Parameters
        ----------
        md:
            Markdown content.

        Returns
        -------
        bytes
            PDF file content.
        """
        # Try pandoc first
        has_pandoc = await self._check_pandoc()
        if has_pandoc:
            with tempfile.NamedTemporaryFile(suffix=".md", mode="w",
                                              encoding="utf-8", delete=False) as f:
                f.write(md)
                tmp_in = f.name

            tmp_out = str(self._save_path("pdf"))
            try:
                await self._run_pandoc(
                    tmp_in, "-f", "markdown", "-o", tmp_out,
                    "--pdf-engine", "weasyprint",
                )
                result = Path(tmp_out).read_bytes()
                return result
            except ConversionError:
                                import logging as _log; _log.getLogger('documents.converter').debug('Suppressed exception', exc_info=True)
            finally:
                Path(tmp_in).unlink(missing_ok=True)

        return self._md_to_pdf_python(md)

    def save(self, data: bytes, path: str | Path | None = None, ext: str = "pdf") -> Path:
        """Save converted document to a file.

        Parameters
        ----------
        data:
            Document content.
        path:
            Output path.
        ext:
            File extension.

        Returns
        -------
        Path
            Saved file path.
        """
        save_to = Path(path) if path else self._save_path(ext)
        save_to.parent.mkdir(parents=True, exist_ok=True)
        save_to.write_bytes(data)
        log.info("Saved converted document → %s (%d bytes)", save_to, len(data))
        return save_to
